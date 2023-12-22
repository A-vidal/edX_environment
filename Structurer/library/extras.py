import docx

# get the html logo of an svg
def get_logo(path: str) -> str:
    # path: where is the logo file
    # output: the html data of the logo
    svg = open(path, "r")
    logo = svg.read()
    svg.close()
    del svg
    return logo

# convert a word content into a html content
def docx_to_html(paragraphs: list) -> str:
    # paragraphs: the word element paragraph atribute
    # output: a html like string
    ret = ""

    for i in paragraphs:
        if i.style.name == "List Paragraph":
            ret += f"<li>{i.text}</li>\n"
        else:
            ret += f"<p>{i.text}</p>\n"

    return ret + "<br>"

class Doc:
    # open and read a word file and return an object with all the data
    def __init__(self, filename: str) -> None:
        # filename: name or path of the word file
        self.filename = filename
        self.word = docx.Document(filename)

        self.intro = self.word.tables[0]
        self.description = self.word.tables[1]

        self.intro_dic = {
            self.intro.cell(0,0).text[:-1]: self.intro.cell(0,1).text,
            self.intro.cell(1,0).text[:-1]: self.intro.cell(1,1).text,
            self.intro.cell(2,0).text[:-1]: self.intro.cell(2,1).text
        }

        self.description_dic = {
            "MC title": self.description.cell(0,1).text,
            "MC long title": self.description.cell(1,1).text,
            
            "ISCED codes1": self.description.cell(2,1).text,
            "EQF level2,3": self.description.cell(2,3).text,
            "ECTS": self.description.cell(2,5).text,
            
            "Suitable for": self.description.cell(3,1).text,
            
            "Background of the proposed micro-credential": docx_to_html(self.description.cell(5,0).paragraphs),
            "Overview of the micro-credential": docx_to_html(self.description.cell(7,0).paragraphs),
            "Learning objectives": docx_to_html(self.description.cell(9,0).paragraphs),
            "Table of contents": docx_to_html(self.description.cell(11,0).paragraphs),
            "Teaching and learning methods": docx_to_html(self.description.cell(13,0).paragraphs),
            "Prerequisites": docx_to_html(self.description.cell(15,0).paragraphs),
            "Assessment methods4": docx_to_html(self.description.cell(17,0).paragraphs)
        }

        self.desc_overview = self.description.cell(7,0).text
